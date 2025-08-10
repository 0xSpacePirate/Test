import os
import sys
import sqlite3
import platform
import warnings

# Suppress warnings from libraries to keep the output clean
warnings.filterwarnings("ignore", category=UserWarning)

def get_system_info():
    """Get system information for troubleshooting."""
    system = platform.system()
    return system

def try_extraction_libraries(filepath):
    """
    Tries a series of Python libraries to extract text from a document.
    'unstructured' is prioritized as it's generally the most effective.
    """
    extracted_content = None
    try:
        from unstructured.partition.auto import partition
        elements = partition(filename=filepath)
        content = "\n\n".join([str(el) for el in elements])
        if content and content.strip():
            extracted_content = content
        else:
            print("✗ 'unstructured' returned empty content.")
    except Exception as e:
        if "pandoc" in str(e).lower():
            print("✗ unstructured ERROR: Pandoc not found. Please install it from pandoc.org.")
        else:
            print(f"✗ unstructured ERROR: {e}")
            
    if not extracted_content:
        print(" -> Trying 'textract' as a fallback...")
        try:
            import textract
            byte_content = textract.process(filepath)
            content = byte_content.decode('utf-8', errors='ignore')
            if content and content.strip():
                extracted_content = content
            else:
                print("✗ 'textract' returned empty content.")
        except Exception as e:
            print(f"✗ textract ERROR: {e}")

    return extracted_content

def manual_dependency_guide():
    """Provides a guide for installing the necessary external programs."""
    print("\n" + "="*60)
    print("AUTOMATIC EXTRACTION FAILED")
    print("="*60)
    print("Please ensure the required dependencies are installed:")
    print("\nOPTION 1 (Recommended): INSTALL PANDOC")
    print("1. Go to https://pandoc.org/installing.html and run the Windows installer.")
    print("2. In your terminal, run: pip install \"unstructured[doc]\"")
    print("\nAfter installation, run this script again.")
    print("="*60)

def extract_doc_text(filepath):
    """
    Extracts text from a single .doc file. Returns None on failure.
    """
    if not os.path.exists(filepath):
        print(f"  -> File not found: {filepath}")
        return None
        
    content = try_extraction_libraries(filepath)

    if content and content.strip():
        return content.replace('\x00', '').strip()
    return None

def create_db(db_path='documents.db'):
    """Creates the database tables needed for searching and tracking indexed files."""
    try:
        conn = sqlite3.connect(db_path)
        c = conn.cursor()
        # Table for full-text search content
        c.execute('''
        CREATE VIRTUAL TABLE IF NOT EXISTS documents USING fts5(
            filename,
            content,
            tokenize = 'porter unicode61'
        );
        ''')
        # Table to track which file paths have already been indexed
        c.execute("CREATE TABLE IF NOT EXISTS indexed_files (path TEXT PRIMARY KEY)")
        conn.commit()
        conn.close()
        print(f"Database '{db_path}' is ready.")
    except sqlite3.Error as e:
        print(f"Database error: {e}")
        sys.exit(1)

def get_indexed_files(db_path):
    """
    --- THIS IS A KEY FUNCTION ---
    Queries the database and returns a set of all file paths that have been processed.
    """
    try:
        conn = sqlite3.connect(db_path)
        c = conn.cursor()
        c.execute("SELECT path FROM indexed_files")
        # Using a set {} gives very fast lookups
        indexed = {row[0] for row in c.fetchall()}
        conn.close()
        return indexed
    except sqlite3.Error as e:
        print(f"Database error checking indexed files: {e}")
        return set()

def insert_file_to_sqlite(filepath, content, db_path='documents.db'):
    """
    --- THIS IS A KEY FUNCTION ---
    Inserts file content and also records the file's path in the 'indexed_files' table.
    """
    if not content or not content.strip():
        print("  -> Content is empty. Skipping database insertion.")
        return False

    try:
        conn = sqlite3.connect(db_path)
        c = conn.cursor()
        # Insert content for searching
        c.execute('INSERT OR REPLACE INTO documents (filename, content) VALUES (?, ?)', (os.path.basename(filepath), content))
        # MARK the file as indexed by inserting its path.
        c.execute('INSERT OR REPLACE INTO indexed_files (path) VALUES (?)', (filepath,))
        conn.commit()
        conn.close()
        return True
    except sqlite3.Error as e:
        print(f"  -> FAILED to insert into database: {e}")
        return False

def search_sqlite(query, db_path='documents.db'):
    """Performs a full-text search on the 'content' column of the database."""
    results = []
    try:
        conn = sqlite3.connect(db_path)
        c = conn.cursor()
        c.execute("SELECT filename FROM documents WHERE documents MATCH ?", (f'"{query}"',))
        results = c.fetchall()
        conn.close()
    except sqlite3.Error as e:
        print(f"Search error: {e}")
    return [filename for (filename,) in results]

def start_search_loop(db_path='documents.db'):
    """Starts the interactive search loop."""
    print("\n" + "="*50)
    print("🔍 ALL DOCUMENTS INDEXED. SEARCH IS READY.")
    print("="*50)
    print("Type any word or phrase (e.g., 'report', 'български').")
    print("Type 'exit' to quit.")
    print("="*50)

    while True:
        q = input("\n🔍 Search: ").strip()
        if not q:
            continue
        if q.lower() == 'exit':
            print("Goodbye!")
            break

        print(f"\nSearching for: '{q}'")
        found_files = search_sqlite(q, db_path)

        if found_files:
            print(f"\n✓ Found '{q}' in the following files:")
            for i, filename in enumerate(found_files, 1):
                print(f"  {i}. {filename}")
        else:
            print("❌ No matching documents found.")

def process_directory(source_directory, db_path):
    """
    Finds all .doc and .docx files, and only processes the ones that are new.
    """
    if not os.path.isdir(source_directory):
        print(f"Error: Directory '{source_directory}' not found.")
        return 0

    print(f"Scanning for documents in '{source_directory}'...")
    
    files_to_process = []
    for root, _, files in os.walk(source_directory):
        for file in files:
            if file.lower().endswith(('.doc', '.docx')):
                # Get the full, absolute path for a unique identifier
                files_to_process.append(os.path.abspath(os.path.join(root, file)))
    
    if not files_to_process:
        print("No .doc or .docx files found in the directory.")
        return 0
        
    # --- LOGIC TO AVOID RE-PROCESSING ---
    # 1. Get the list of files already in the database.
    already_indexed = get_indexed_files(db_path)
    
    # 2. Subtract the 'already_indexed' set from the 'files_to_process' list.
    new_files = [f for f in files_to_process if f not in already_indexed]
    # --- END OF LOGIC ---

    if not new_files:
        print(f"All {len(files_to_process)} documents are already indexed. Nothing new to add.")
        return len(files_to_process)

    print(f"Found {len(new_files)} new documents to process.")
    
    successful_imports = 0
    failed_imports = 0
    
    for i, filepath in enumerate(new_files):
        print("\n" + "-"*60)
        print(f"Processing file {i+1} of {len(new_files)}: {os.path.basename(filepath)}")
        
        content = extract_doc_text(filepath)
        
        if content:
            if insert_file_to_sqlite(filepath, content, db_path):
                print(f"  -> Successfully indexed.")
                successful_imports += 1
            else:
                failed_imports += 1
        else:
            print(f"  -> FAILED to extract text.")
            failed_imports += 1
    
    print("\n" + "="*60)
    print("BATCH PROCESSING COMPLETE")
    print(f"Successfully Indexed: {successful_imports}")
    print(f"Failed: {failed_imports}")
    print("="*60)
    
    return len(get_indexed_files(db_path))

def main():
    """Main function to set up and run the application."""
    print("Universal .doc File Text Extraction and Search (Batch Mode)")
    print("=" * 60)
    get_system_info()
    
    source_directory = 'documents' # Changed from 'test' for clarity
    db_path = 'document_search.db'
    
    if not os.path.exists(source_directory):
        print(f"Creating a sample '{source_directory}/' folder for you...")
        os.makedirs(source_directory)
        try:
            from docx import Document
            doc1 = Document()
            doc1.add_paragraph("This is the first document. It contains the word 'alpha'.")
            doc1.add_paragraph("Това е първият документ на български език.")
            doc1.save(os.path.join(source_directory, 'report_alpha.doc'))
            doc2 = Document()
            doc2.add_paragraph("This is the second file, mentioning 'beta'.")
            doc2.save(os.path.join(source_directory, 'summary_beta.docx'))
            print("Created two sample documents inside the folder.")
        except ImportError:
            print(f"Could not create dummy files: 'python-docx' not installed. Please manually add .doc files to the '{source_directory}' folder.")
        except Exception as e:
            print(f"An error occurred while creating dummy files: {e}")
            
    create_db(db_path)

    total_indexed_docs = process_directory(source_directory, db_path)

    if total_indexed_docs > 0:
        start_search_loop(db_path)
    else:
        print("\nNo documents could be indexed. The search cannot start.")
        manual_dependency_guide()
        sys.exit(1)


if __name__ == "__main__":
    main()